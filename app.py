from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, make_response
from flask_mail import Mail, Message
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from config import Config
from models import db, User, Submission
import os

app = Flask(__name__)
app.config.from_object(Config)

# Initialize extensions
db.init_app(app)
mail = Mail(app)
login_manager = LoginManager(app)
login_manager.login_view = 'admin_login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Context processor to inject current year into templates
@app.context_processor
def inject_year():
    from datetime import datetime
    return {'year': datetime.now().year}

# --- Public Routes ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/services')
def services():
    return render_template('services.html')

@app.route('/pricing')
def pricing():
    from models import PricingPlan
    import json
    
    plans = PricingPlan.query.order_by(PricingPlan.id).all()
    
    # Convert features JSON string to list for template
    for plan in plans:
        if plan.features:
            plan.features_list = json.loads(plan.features)
        else:
            plan.features_list = []
    
    return render_template('pricing.html', plans=plans)


@app.route('/contact', methods=['GET', 'POST'])
def contact():
    if request.method == 'POST':
        try:
            # Extract form data
            full_name = request.form.get('full_name')
            business_name = request.form.get('business_name')
            email = request.form.get('email')
            whatsapp_number = request.form.get('whatsapp_number')
            country = request.form.get('country')
            message = request.form.get('message')
            plan_selected = request.form.get('plan_selected')
            
            # Create new submission
            new_submission = Submission(
                full_name=full_name,
                business_name=business_name,
                email=email,
                whatsapp_number=whatsapp_number,
                country=country,
                message=message,
                plan_selected=plan_selected
            )
            
            db.session.add(new_submission)
            db.session.commit()
            
            # Send email notification
            try:
                msg = Message(
                    subject="New WhatsFlow Client Submission",
                    recipients=[app.config['ADMIN_EMAIL']],
                    html=render_template('email/new_submission.html', submission=new_submission)
                )
                mail.send(msg)
            except Exception as e:
                print(f"Email sending failed: {e}")
                # Continue even if email fails, but log it
            
            return redirect(url_for('success'))
            
        except Exception as e:
            print(f"Submission failed: {e}")
            flash('Something went wrong. Please try again.', 'error')
            return redirect(url_for('contact'))
            
    from models import PricingPlan
    plans = PricingPlan.query.order_by(PricingPlan.id).all()
    return render_template('contact.html', plans=plans)

@app.route('/success')
def success():
    return render_template('success.html')

# --- Admin Routes ---

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if current_user.is_authenticated:
        return redirect(url_for('admin_dashboard'))
        
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        user = User.query.filter_by(email=email).first()
        
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid email or password', 'error')
            
    return render_template('admin/login.html')

@app.route('/admin/logout')
@login_required
def admin_logout():
    logout_user()
    return redirect(url_for('admin_login'))

@app.route('/admin')
@login_required
def admin_dashboard():
    # Dashboard stats
    total_submissions = Submission.query.count()
    pending_submissions = Submission.query.filter_by(status='pending').count()
    contacted_submissions = Submission.query.filter_by(status='contacted').count()
    
    # Recent submissions
    recent_submissions = Submission.query.order_by(Submission.created_at.desc()).limit(5).all()
    
    return render_template('admin/dashboard.html', 
                          total=total_submissions,
                          pending=pending_submissions,
                          contacted=contacted_submissions,
                          recent=recent_submissions)

from flask import send_file, make_response
import csv
import json
import io
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

@app.route('/admin/submissions')
@login_required
def admin_submissions():
    submissions = Submission.query.order_by(Submission.created_at.desc()).all()
    return render_template('admin/submissions.html', submissions=submissions)

@app.route('/admin/export/<format>')
@login_required
def export_data(format):
    try:
        import tempfile
        import os
        from flask import send_from_directory
        
        submissions = Submission.query.order_by(Submission.created_at.desc()).all()
        
        # Create temp directory if it doesn't exist
        temp_dir = os.path.join(app.root_path, 'temp_exports')
        os.makedirs(temp_dir, exist_ok=True)
        
        if format == 'csv':
            filename = 'submissions.csv'
            filepath = os.path.join(temp_dir, filename)
            
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['Date', 'Name', 'Business', 'Email', 'WhatsApp', 'Country', 'Plan', 'Message', 'Status'])
                for sub in submissions:
                    writer.writerow([
                        sub.created_at.strftime('%Y-%m-%d %H:%M'),
                        sub.full_name,
                        sub.business_name,
                        sub.email,
                        sub.whatsapp_number,
                        sub.country,
                        sub.plan_selected,
                        sub.message,
                        sub.status
                    ])
            
            return send_from_directory(temp_dir, filename, as_attachment=True)
            
        elif format == 'json':
            filename = 'submissions.json'
            filepath = os.path.join(temp_dir, filename)
            
            data = []
            for sub in submissions:
                data.append({
                    'date': sub.created_at.strftime('%Y-%m-%d %H:%M'),
                    'name': sub.full_name,
                    'business': sub.business_name,
                    'email': sub.email,
                    'whatsapp': sub.whatsapp_number,
                    'country': sub.country,
                    'plan': sub.plan_selected,
                    'message': sub.message,
                    'status': sub.status
                })
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=4)
            
            return send_from_directory(temp_dir, filename, as_attachment=True)
            
        elif format == 'excel':
            filename = 'submissions.xlsx'
            filepath = os.path.join(temp_dir, filename)
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Submissions"
            ws.append(['Date', 'Name', 'Business', 'Email', 'WhatsApp', 'Country', 'Plan', 'Message', 'Status'])
            for sub in submissions:
                ws.append([
                    sub.created_at.strftime('%Y-%m-%d %H:%M'),
                    sub.full_name,
                    sub.business_name,
                    sub.email,
                    sub.whatsapp_number,
                    sub.country,
                    sub.plan_selected,
                    sub.message,
                    sub.status
                ])
            wb.save(filepath)
            
            return send_from_directory(temp_dir, filename, as_attachment=True)
            
        elif format == 'pdf':
            filename = 'submissions.pdf'
            filepath = os.path.join(temp_dir, filename)
            
            doc = SimpleDocTemplate(filepath, pagesize=landscape(letter))
            elements = []
            styles = getSampleStyleSheet()
            
            elements.append(Paragraph("WhatsFlow Submissions", styles['Title']))
            
            data = [['Date', 'Name', 'Business', 'Email', 'Plan', 'Status']]
            for sub in submissions:
                data.append([
                    sub.created_at.strftime('%Y-%m-%d'),
                    sub.full_name,
                    sub.business_name,
                    sub.email,
                    sub.plan_selected,
                    sub.status
                ])
                
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(table)
            doc.build(elements)
            
            return send_from_directory(temp_dir, filename, as_attachment=True)
            
        elif format == 'word':
            filename = 'submissions.docx'
            filepath = os.path.join(temp_dir, filename)
            
            doc = Document()
            doc.add_heading('WhatsFlow Submissions', 0)
            
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Date'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Business'
            hdr_cells[3].text = 'Email'
            hdr_cells[4].text = 'Plan'
            hdr_cells[5].text = 'Status'
            
            for sub in submissions:
                row_cells = table.add_row().cells
                row_cells[0].text = sub.created_at.strftime('%Y-%m-%d')
                row_cells[1].text = sub.full_name
                row_cells[2].text = sub.business_name
                row_cells[3].text = sub.email
                row_cells[4].text = sub.plan_selected
                row_cells[5].text = sub.status
                
            doc.save(filepath)
            
            return send_from_directory(temp_dir, filename, as_attachment=True)
            
        return redirect(url_for('admin_submissions'))
        
    except Exception as e:
        print(f"Export failed: {e}")
        import traceback
        traceback.print_exc()
        flash(f"Export failed: {str(e)}", 'error')
        return redirect(url_for('admin_submissions'))

@app.route('/admin/submission/<int:id>/update-status', methods=['POST'])
@login_required
def update_submission_status(id):
    submission = Submission.query.get_or_404(id)
    new_status = request.form.get('status')
    if new_status:
        submission.status = new_status
        db.session.commit()
        flash(f'Status updated to {new_status}', 'success')
    return redirect(url_for('admin_submissions'))

@app.route('/admin/settings')
@login_required
def admin_settings():
    return render_template('admin/settings.html')

@app.route('/admin/pricing')
@login_required
def admin_pricing():
    from models import PricingPlan
    plans = PricingPlan.query.order_by(PricingPlan.id).all()
    return render_template('admin/pricing.html', plans=plans)

@app.route('/admin/pricing/update/<int:plan_id>', methods=['POST'])
@login_required
def update_pricing(plan_id):
    from models import PricingPlan
    plan = PricingPlan.query.get_or_404(plan_id)
    
    try:
        base_price = float(request.form.get('base_price', plan.base_price))
        discount_percent = int(request.form.get('discount_percent', 0))
        
        plan.base_price = base_price
        plan.apply_discount(discount_percent)
        
        db.session.commit()
        flash(f'{plan.plan_name} plan updated successfully!', 'success')
    except Exception as e:
        flash(f'Error updating pricing: {str(e)}', 'error')
    
    return redirect(url_for('admin_pricing'))

@app.route('/admin/settings/change-email', methods=['POST'])
@login_required
def admin_change_email():
    new_email = request.form.get('new_email')
    current_password = request.form.get('current_password')
    
    # Validate current password
    if not current_user.check_password(current_password):
        flash('Current password is incorrect.', 'error')
        return redirect(url_for('admin_settings'))
    
    # Validate new email
    if not new_email or '@' not in new_email:
        flash('Please enter a valid email address.', 'error')
        return redirect(url_for('admin_settings'))
    
    # Check if email is already in use
    existing_user = User.query.filter_by(email=new_email).first()
    if existing_user and existing_user.id != current_user.id:
        flash('This email is already in use.', 'error')
        return redirect(url_for('admin_settings'))
    
    # Update email
    current_user.email = new_email
    db.session.commit()
    
    # Log out user
    logout_user()
    flash('Email updated successfully. Please log in with your new email.', 'success')
    return redirect(url_for('admin_login'))

@app.route('/admin/settings/change-password', methods=['POST'])
@login_required
def admin_change_password():
    current_password = request.form.get('current_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    
    # Validate current password
    if not current_user.check_password(current_password):
        flash('Current password is incorrect.', 'error')
        return redirect(url_for('admin_settings'))
    
    # Validate new password
    if not new_password or len(new_password) < 8:
        flash('New password must be at least 8 characters long.', 'error')
        return redirect(url_for('admin_settings'))
    
    # Check passwords match
    if new_password != confirm_password:
        flash('New passwords do not match.', 'error')
        return redirect(url_for('admin_settings'))
    
    # Update password
    current_user.set_password(new_password)
    db.session.commit()
    
    # Log out user
    logout_user()
    flash('Password updated successfully. Please log in with your new password.', 'success')
    return redirect(url_for('admin_login'))


# --- CLI Commands ---
@app.cli.command("init-db")
def init_db_command():
    """Clear existing data and create new tables."""
    db.create_all()
    print("Initialized the database.")

@app.cli.command("create-admin")
def create_admin_command():
    """Create a default admin user."""
    email = input("Enter admin email: ")
    password = input("Enter admin password: ")
    
    if User.query.filter_by(email=email).first():
        print("User already exists.")
        return
        
    user = User(email=email, is_admin=True)
    user.set_password(password)
    db.session.add(user)
    db.session.commit()
    print(f"Admin user {email} created successfully.")

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
